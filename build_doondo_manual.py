"""
Build the Doondo end-user manual as a .docx file.

End-user audience — plain English, step-by-step, screen-by-screen.
Covers both the Seeker (worker) and Employer flows from account
creation through every feature in the app.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import os
# Resolve output path so this script works both via the Cowork file
# tools (/Users/...) and from the sandboxed bash (/sessions/...).
_candidates = [
    "/Users/_itsshree/Doondo V2",
    "/sessions/festive-pensive-lovelace/mnt/Doondo V2",
]
_base = next((c for c in _candidates if os.path.isdir(c)), _candidates[0])
OUT = os.path.join(_base, "Doondo_User_Manual.docx")

# ─── Setup ───────────────────────────────────────────────────────────────────

doc = Document()

# Page setup — US Letter, 1" margins
for section in doc.sections:
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Default font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.25

# Tune heading styles
def tune_heading(name, size, color=(0x1F, 0x2A, 0x44), space_before=14, space_after=6, bold=True):
    h = doc.styles[name]
    h.font.name = "Calibri"
    h.font.size = Pt(size)
    h.font.bold = bold
    h.font.color.rgb = RGBColor(*color)
    h.paragraph_format.space_before = Pt(space_before)
    h.paragraph_format.space_after = Pt(space_after)
    h.paragraph_format.keep_with_next = True

tune_heading("Title", 28, color=(0x25, 0x63, 0xEB), space_before=0, space_after=4)
tune_heading("Heading 1", 22, color=(0x25, 0x63, 0xEB), space_before=20, space_after=8)
tune_heading("Heading 2", 16, color=(0x1F, 0x2A, 0x44), space_before=14, space_after=4)
tune_heading("Heading 3", 13, color=(0x2D, 0x3B, 0x55), space_before=10, space_after=2)

# ─── Helpers ─────────────────────────────────────────────────────────────────

def p(text="", bold=False, italic=False, size=11, color=None, indent=None, after=None):
    para = doc.add_paragraph()
    if indent is not None:
        para.paragraph_format.left_indent = Inches(indent)
    if after is not None:
        para.paragraph_format.space_after = Pt(after)
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor(*color)
    return para

def steps(items, start=1):
    """Numbered list with inline numbering — restarts at 1 every call.

    Avoids python-docx's shared List Number counter, which would
    continue numbers across separate procedural sections.
    """
    for i, item in enumerate(items, start=start):
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.45)
        para.paragraph_format.first_line_indent = Inches(-0.3)
        para.paragraph_format.space_after = Pt(3)
        num_run = para.add_run(f"{i}.\t")
        num_run.bold = True
        num_run.font.size = Pt(11)
        run = para.add_run(item)
        run.font.size = Pt(11)
        # Tab stop so the body lines up cleanly after "10."
        para.paragraph_format.tab_stops.add_tab_stop(Inches(0.45))

def bullets(items):
    """Bulleted list."""
    for item in items:
        para = doc.add_paragraph(style="List Bullet")
        para.paragraph_format.left_indent = Inches(0.35)
        para.paragraph_format.space_after = Pt(3)
        # Item may be a tuple (label, body) — render with bold label
        if isinstance(item, tuple):
            label, body = item
            r1 = para.add_run(label)
            r1.bold = True
            r1.font.size = Pt(11)
            r2 = para.add_run(" — " + body)
            r2.font.size = Pt(11)
        else:
            run = para.add_run(item)
            run.font.size = Pt(11)

def _insert_into_pPr(pPr, element, after_tags=(), before_tags=()):
    """Insert `element` into pPr respecting OOXML schema order.

    after_tags  — list of tag names that must come BEFORE element (skip past them).
    before_tags — list of tag names that must come AFTER element (insert before them).
    """
    # Find the right index: just before the first child in before_tags,
    # or right after the last child in after_tags, or at the start.
    insert_at = 0
    for i, child in enumerate(list(pPr)):
        tag = child.tag.split("}", 1)[1] if "}" in child.tag else child.tag
        if tag in after_tags:
            insert_at = i + 1
        elif tag in before_tags:
            insert_at = i
            break
    pPr.insert(insert_at, element)


def section_break():
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "DBDFE6")
    pBdr.append(bottom)
    # pBdr goes after pStyle/numPr, before shd/tabs/spacing/ind/jc/rPr
    _insert_into_pPr(
        pPr, pBdr,
        after_tags=("pStyle", "numPr", "suppressLineNumbers"),
        before_tags=("shd", "tabs", "spacing", "ind", "contextualSpacing",
                     "jc", "outlineLvl", "rPr"),
    )

def tip(label, body):
    """A highlighted callout — light-yellow background paragraph."""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.1)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(8)
    pPr = para._p.get_or_add_pPr()
    # Left border — pBdr must come BEFORE shd in schema order
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "F59E0B")
    pBdr.append(left)
    _insert_into_pPr(
        pPr, pBdr,
        after_tags=("pStyle", "numPr", "suppressLineNumbers"),
        before_tags=("shd", "tabs", "spacing", "ind", "contextualSpacing",
                     "jc", "outlineLvl", "rPr"),
    )
    # Background shading — goes AFTER pBdr, BEFORE spacing/ind/jc/rPr
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "FEF7E0")
    _insert_into_pPr(
        pPr, shd,
        after_tags=("pStyle", "numPr", "suppressLineNumbers", "pBdr"),
        before_tags=("tabs", "spacing", "ind", "contextualSpacing",
                     "jc", "outlineLvl", "rPr"),
    )
    r1 = para.add_run(label + " ")
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = RGBColor(0xB4, 0x53, 0x09)
    r2 = para.add_run(body)
    r2.font.size = Pt(11)

def h1(text):
    doc.add_heading(text, level=1)

def h2(text):
    doc.add_heading(text, level=2)

def h3(text):
    doc.add_heading(text, level=3)

def page_break():
    doc.add_page_break()

# ─── Title page ──────────────────────────────────────────────────────────────

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(120)
r = title.add_run("Doondo")
r.font.name = "Calibri"
r.font.size = Pt(54)
r.bold = True
r.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_after = Pt(40)
r = sub.add_run("User Manual")
r.font.size = Pt(22)
r.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)

tag = doc.add_paragraph()
tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tag.add_run("Hiring made simple — for workers and employers.")
r.italic = True
r.font.size = Pt(13)
r.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

box_p = doc.add_paragraph()
box_p.paragraph_format.space_before = Pt(80)
box_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
box_p.add_run(
    "This guide walks you through every screen in the Doondo mobile app, "
    "from creating your account to landing a job (or hiring one). "
    "Use the table of contents on the next page to jump to the section you need."
).font.size = Pt(11)

doc.add_page_break()

# ─── Table of contents (manual) ──────────────────────────────────────────────

h1("What's in this guide")

p("Part 1 — Getting started (for everyone)", bold=True, size=12)
bullets([
    "Welcome to Doondo",
    "Your first launch",
    "Picking your role: worker or employer",
    "Creating an account",
    "Signing in",
    "Forgot your password",
])

p("Part 2 — For workers (job seekers)", bold=True, size=12)
bullets([
    "Your home dashboard",
    "Voice search",
    "Browsing jobs by category",
    "Job details and applying",
    "Cover letters",
    "My Applications",
    "My Jobs (saved + applied)",
    "Chat with employers",
    "Your profile",
    "Editing profile sections",
    "Phone verification and selfie",
    "Resume Builder",
    "Resume Preview and sharing",
    "Job Alerts",
    "Notifications",
    "My Earnings (cash income diary)",
    "Ratings and reviews",
    "Download Center",
    "Settings (language, theme, notifications)",
    "Safety SOS",
    "Skill Tests",
    "Courses and training",
    "Available now beacon",
    "Referrals — share and earn",
    "Group hire — apply with teammates",
])

p("Part 3 — For employers", bold=True, size=12)
bullets([
    "Your job posts",
    "Posting a new job",
    "Reviewing applicants",
    "Applicant details",
    "Hiring and rating workers",
    "Chat with workers",
    "Your employer profile",
    "Available workers (live nearby)",
    "Verifying a worker's photos",
])

p("Part 4 — Help and reference", bold=True, size=12)
bullets([
    "Staying safe",
    "Switching languages",
    "Account help",
    "Common questions",
    "Contacting support",
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# PART 1 — GETTING STARTED
# ════════════════════════════════════════════════════════════════════════════

h1("Part 1 — Getting started")

h2("Welcome to Doondo")
p(
    "Doondo is a hiring app built for India. It connects workers — from "
    "drivers, cooks, electricians, and delivery riders to designers, "
    "developers, and accountants — directly with employers near them. "
    "There are no middlemen and no posting fees."
)
p(
    "Doondo works for two kinds of people:"
)
bullets([
    ("Workers (job seekers)", "find work, apply with one tap, get paid, build a verified reputation."),
    ("Employers", "post jobs in minutes, see who's available nearby right now, and hire the right person."),
])
tip("Good to know:", "The same app handles both blue-collar and white-collar work. Whether you're hiring a plumber or a UX designer, the flow is the same.")

h2("Your first launch")
p("The first time you open Doondo, you'll see a short three-screen walkthrough explaining what the app does. Swipe through it or tap Skip in the top-right to jump straight to the role picker.")

h2("Picking your role")
p("After the walkthrough, Doondo asks whether you're here to find work or to hire. This is the role picker — a full-screen 3D animated card.")
steps([
    "Tap I'm looking for work if you want to find jobs.",
    "Tap I'm hiring if you want to post jobs and find workers.",
    "If you're an employer, you'll also be asked roughly how many people you usually hire (Just me / Small team / Larger). This helps Doondo tune what you see — it's not a commitment.",
])
tip("You can switch later:", "Your role isn't locked. From Settings → Account you can re-register with a different role if your needs change.")

h2("Creating an account")
p("Sign-up is short — Doondo only asks for what's needed to keep your account safe and to let employers contact you.")
steps([
    "Enter your full name. Workers: this is the name employers will see. Employers: this is your business or personal name.",
    "Enter the phone number where you want to receive calls and password-reset codes. Workers — be sure this is the number a future employer can call.",
    "Enter your email and a password (at least 8 characters).",
    "If you're setting the account up for someone else (common — family members helping a worker who can't use a phone alone), turn on Setting this up for someone else? The phone field then asks for the worker's number instead of yours.",
    "Tap Create account. Doondo creates your profile and signs you in.",
])
tip("Why phone is required:", "If you ever forget your password, the only way Doondo can help you back in is via a code sent to your phone. Make sure it's a number you actually use.")

h2("Signing in")
p("If you already have an account:")
steps([
    "Open Doondo.",
    "On the welcome screen, tap I already have an account.",
    "Enter your email and password.",
    "Tap Sign in.",
])

h2("Forgot your password")
p("Forgot your password? You can reset it in three quick steps:")
steps([
    "On the sign-in screen, tap Forgot password?",
    "Enter the phone number on your account. Doondo sends a 6-digit code by SMS.",
    "Enter the code on the next screen. The code is valid for a few minutes — if it expires, request a new one.",
    "Pick a new password and confirm. You'll be signed in automatically.",
])
tip("Don't have a phone on file?", "Older accounts (created before phone was required) can add a recovery phone any time from Profile → Settings → Account. Doondo will ask you to verify it with a code, just like signup.")

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# PART 2 — WORKERS
# ════════════════════════════════════════════════════════════════════════════

h1("Part 2 — For workers (job seekers)")

p("The seeker side of Doondo has four tabs at the bottom of the screen: Home, Jobs, Chat, and Profile. We'll walk through each, then cover every modal you can reach from them.")

h2("Your home dashboard")
p("Home is the first thing you see after signing in. It's organised around three modes that match how Indian workers actually look for jobs:")
bullets([
    ("Today", "Same-day work that starts within hours. Tap the Today tab to see only urgent, ASAP postings."),
    ("This week", "Short-term work — daily and weekly contracts."),
    ("Career", "Full-time and longer-term roles. This is the default."),
])
p("Switch modes by tapping the segmented control at the top of Home. Your choice is remembered so the next time you open the app, you land back where you were.")

h3("What's on the home screen")
bullets([
    ("Voice search card", "A big blue card at the top that lets you describe what you want in your own language."),
    ("Job categories", "Tap a tile (Driving, Cooking, Electrical, etc.) to filter the jobs list to that trade."),
    ("Nearby jobs", "A preview of jobs sorted by distance. Tap View all to open the full list."),
    ("Available now toggle", "Workers only — a switch that broadcasts \"I'm available right now\" to nearby employers. More on this below."),
])

h2("Voice search")
p("Voice search is the fastest way to find work if typing is awkward or you'd rather speak in Hindi, Kannada, Tamil, or Telugu.")
steps([
    "Tap the voice card on Home, or the microphone button at the centre of the bottom navigation.",
    "Hold the big blue mic button and speak — for example, \"painting work near Jayanagar\" or \"daily wage construction\".",
    "Let go when you finish. Doondo transcribes what you said.",
    "Tap Search. The jobs list opens, filtered to match what you described.",
])
tip("Privacy:", "Voice recordings are processed for transcription and then discarded. Doondo doesn't keep audio of what you said.")

h2("Browsing jobs by category")
p("From Home, tap a category tile to filter jobs. The Jobs tab also has filters along the top:")
bullets([
    ("Distance", "1 km, 5 km, 10 km, or any distance."),
    ("Job type", "Full-time, part-time, daily, weekly, or contract."),
    ("Work style", "Onsite, hybrid, or remote — relevant for white-collar work."),
    ("Urgent only", "Hide non-urgent posts."),
])
p("Tap any filter chip to change it. Doondo updates the list immediately. To clear all filters, scroll to the top and tap Clear.")

h2("Job details and applying")
p("Tap any job card to open the full job detail screen.")
p("You'll see:")
bullets([
    "The employer's business name, photo, and verified badge (if they've completed verification).",
    "The job title, description, and any audio description (tap play to hear the employer describe the job in their own words).",
    "Daily pay or salary range — Doondo shows whether the rate is per day, per shift, per week, or per month.",
    "Distance from your saved location.",
    "Skills required and whether the employer is open to teaching.",
    "Photos the employer posted of the worksite or work to be done.",
])
h3("Applying")
p("The Apply button is at the bottom of the screen and stays visible as you scroll.")
steps([
    "Tap Apply Now.",
    "Optionally write a short cover letter (1–2 sentences about why you'd be a good fit). You can leave it blank if you prefer.",
    "If the job is a group hire (the employer wants to hire multiple workers together), you can add your teammates' names and phone numbers — up to four of them. Doondo treats your application as a team.",
    "Tap Submit. The Apply button changes to Applied ✓ and the employer is notified.",
])
h3("Today-mode CTAs")
p("If you opened the job from the Today tab, the buttons are different — they're tuned for same-day work:")
bullets([
    ("I'm interested", "A one-tap signal to the employer that you're free now. They'll see you in their \"interested\" list and can call you directly."),
    ("Call employer", "Available after you express interest — opens your phone dialer. Doondo doesn't share your number; the employer sees it only after they accept your interest."),
])
h3("Sharing a job")
p("Tap the share icon (top-right of the job detail screen) to share the job with a friend over WhatsApp or any messaging app. If your friend signs up and gets hired through your link, you earn a referral bonus — see Referrals below.")

h2("Cover letters")
p("On any white-collar or career-mode job, you can attach a short cover letter when you apply. Keep it to 1–3 sentences — employers read dozens of applications and a tight pitch wins.")
tip("Pro tip:", "Mention one specific thing from the job posting — the location, the trade, or the employer's name. It shows you read the post.")

h2("My Applications")
p("From your Profile, tap My Applications to see every job you've applied to. Each row shows:")
bullets([
    "Job title and employer name",
    "Date you applied",
    "Current status — Applied, Shortlisted, Hired, or Declined",
    "A chat shortcut if the employer has opened a conversation with you",
])
p("Tap any application to re-open the job detail screen and message the employer.")

h2("My Jobs (saved + applied)")
p("My Jobs combines two lists into one place:")
bullets([
    ("Saved", "Jobs you bookmarked but haven't applied to yet. To save a job, tap the bookmark icon on a job card or the job detail screen."),
    ("Applied", "Same list as My Applications, shown here for convenience."),
])

h2("Chat with employers")
p("The Chat tab lists every conversation you have with an employer. Tap a conversation to open it.")
p("Inside a conversation you can:")
bullets([
    "Send text messages.",
    "Tap the microphone to record a short voice note (up to 60 seconds).",
    "Tap the camera to share a photo (a work sample, ID, or a picture of the worksite).",
    "Tap the video icon to share a short video clip.",
])
tip("Starting a chat:", "Workers don't start chats — employers do, after they shortlist you. If you want to message an employer first, apply for the job and the employer will see your application; they can then open the conversation.")

h2("Your profile")
p("The Profile tab is your home base. Top to bottom, here's what's on it:")
bullets([
    ("Photo, name, and trade", "Your headline. Tap the pen icon to edit any of these."),
    ("Verified badge", "A blue check if you've completed phone and selfie verification. Verified workers get more responses."),
    ("Rating", "An average star score from past employers, with the number of ratings."),
    ("Expected salary", "The pay you're looking for. Employers see this on your profile."),
    ("Skills / trades", "The trades you can work — picked from a 30+ list covering blue and white collar."),
    ("Work photos", "Pictures of jobs you've done. Employers can verify these photos after they hire you — verified photos show a green ✓ badge."),
    ("Menu", "Quick links to My Applications, My Jobs, Earnings, Ratings, Job Alerts, Training, Skill tests, Referrals, Downloads, Settings, and Edit profile."),
])

h2("Editing your profile")
p("Tap Edit Profile Details on your profile, or any pen icon on the profile screen. The edit screen has sections you can update independently:")
bullets([
    ("Basics", "Name, bio, date of birth."),
    ("Location", "Where you're based and willing to work."),
    ("Skills", "The trades you can do."),
    ("Preferences", "Job type (full-time, daily, etc.) and work style (onsite, remote, hybrid)."),
    ("Resume", "Your work history — past jobs, durations, and photos."),
])
p("Employer accounts also have Business basics and Business location sections (covered later).")

h2("Phone verification and selfie")
p("Verification gets you a blue check on your profile, which signals to employers that you're a real person at the phone number you list. Verified workers get noticeably more responses.")
steps([
    "From your profile, tap the Verify your account banner, or go to Settings → Account → Verify.",
    "Doondo sends a 6-digit code to your phone. Enter it on the next screen.",
    "Take a selfie holding your phone. This is just to confirm you're a real person — Doondo doesn't share the photo with employers and doesn't run face matching.",
    "You'll see the blue check on your profile within seconds.",
])

h2("Resume Builder")
p("Resume Builder turns your work history into a clean one-page resume you can share with any employer — even outside Doondo.")
steps([
    "From Profile, tap Edit Profile Details → Resume → Build your resume.",
    "Add 1–5 past jobs. For each, enter the role, employer, dates, and a one-line description. You can attach up to 3 photos per job (worksite, finished work, ID).",
    "Add up to 3 education entries if relevant.",
    "Tap Save. Doondo generates the resume.",
])
p("You can re-open the builder any time to add more jobs or update details.")

h2("Resume Preview and sharing")
p("From Profile, tap Edit Profile Details → Resume → Preview your resume.")
p("The preview shows exactly what an employer or anyone you share with will see.")
bullets([
    ("Share as PDF", "Generates a polished PDF of your resume and opens the share sheet so you can send it via WhatsApp, email, or any other app."),
    ("Verified work photos", "If an employer has verified one of your work photos, it shows a green ✓ Verified by employer badge on the resume — a powerful trust signal."),
    ("Tested trades", "Skill tests you've passed show up as amber TESTED · 5/5 pills."),
])

h2("Job Alerts")
p("Job Alerts let Doondo notify you when a matching job is posted, so you don't have to keep checking the app.")
steps([
    "From Profile, tap Job Alerts.",
    "Tap + New alert.",
    "Give it a name (e.g., \"Driver jobs in Indiranagar\").",
    "Set the keyword, city, job types, and urgent-only toggle.",
    "Tap Save. Doondo will push a notification when a job matching your alert is posted.",
])
p("From the Job Alerts list you can:")
bullets([
    "Tap an alert to edit it.",
    "Toggle it on/off without deleting.",
    "Swipe left to delete.",
])
tip("Suggested alerts:", "Doondo suggests alerts based on your resume — if you list \"electrician\" as a skill, you'll see a one-tap suggestion to create an electrician alert in your city.")

h2("Notifications")
p("The bell icon at the top-right of every tab opens your Notifications feed.")
bullets([
    "New job matches",
    "Application updates (shortlisted, hired, declined)",
    "Messages from employers",
    "Ratings you've received",
    "Referral bonuses earned",
])
p("Tap any notification to jump to the screen it refers to. Tap Clear all to mark everything as read.")

h2("My Earnings (cash income diary)")
p("If you're paid in cash and want to keep your own records, the Earnings ledger gives you a simple income diary.")
steps([
    "From Profile, tap My Earnings.",
    "Tap + Add entry.",
    "Enter the amount, who paid you, what for, and the date.",
    "Tap Save.",
])
p("You'll see running totals for the current week, month, and year. Tap any entry to edit or delete it. The ledger is stored on your device and synced to your account — only you can see it.")

h2("Ratings and reviews")
p("After every hire, employers can rate your work from 1 to 5 stars and write a short review. You can see them all in one place:")
steps([
    "From Profile, tap Ratings & Reviews.",
    "Browse each rating — the star score, the comment, the employer's name, and the job it was for.",
])
p("You can also leave ratings — at the end of every job you complete, Doondo prompts you to rate the employer (1–5 stars + optional comment). Honest ratings help other workers know who's a good employer.")

h2("Download Center")
p("Download Center stores job posts you saved for offline viewing — useful when you have a weak signal but want to re-read a job description.")
steps([
    "On any job detail screen, tap the download icon (top-right, near share).",
    "The job is saved to your device for 30 days.",
    "From Profile → Download Center, browse all saved jobs even without a connection.",
])

h2("Settings")
p("From Profile, tap Settings. The screen has five sections:")
h3("Language")
p("Doondo supports English, Hindi (हिन्दी), Kannada (ಕನ್ನಡ), Tamil (தமிழ்), and Telugu (తెలుగు). Tap a language to switch — the change applies immediately across the whole app.")
h3("Notifications")
p("A master toggle for push notifications. Turn off if you don't want any pings from Doondo. Granular per-type toggles (jobs vs. chat vs. ratings) will arrive in a future update.")
h3("Appearance")
p("Light, Dark, or Use system setting. The worker side stays in its blue theme regardless — this setting affects employer screens and shared modals like the Settings page itself.")
h3("Safety")
p("Quick access to Safety SOS setup (see below).")
h3("Account")
bullets([
    "Email — what's on file (tap to request a change, which requires verification).",
    "Sign out — signs you out on this device. Your account stays intact.",
    "Delete account — permanent. Erases your applications, ratings, and chats. Currently routed through support@doondo.app while we finish the in-app delete flow.",
])

h2("Safety SOS")
p("Safety SOS is a one-tap way to alert a trusted contact if you ever feel unsafe at a worksite.")
steps([
    "From Settings → Safety, or from Profile → Safety SOS, tap Set up emergency contact.",
    "Enter the name and phone number of someone you trust — a family member, friend, or anyone who can come help.",
    "Tap Save. The contact is stored only on your device, never on Doondo's servers.",
])
p("To trigger an SOS:")
steps([
    "From any screen, tap the SOS button in the bottom-right (visible on Home and Jobs).",
    "Confirm — the button asks twice to prevent accidental triggers.",
    "Doondo sends an SMS to your contact with your name and current location, and opens your phone dialer with their number pre-filled.",
])
tip("Safety SOS is not 112:", "It's a personal alert to someone you know. For police or medical emergencies, dial 112 directly.")

h2("Skill Tests")
p("Skill Tests let you prove your trade with a short quiz. Pass once and a tested badge appears on your resume.")
steps([
    "From Profile, tap Skill tests.",
    "Pick a trade from the catalogue (Electrician, Cook, Delivery rider — more added regularly).",
    "Tap Start test. You'll see 5 multiple-choice questions, each with 4 options.",
    "Answer all 5. You need at least 4 correct to pass.",
    "If you pass, the badge is added to your profile immediately. If you fail, you can retake the test in 24 hours.",
])
tip("No retakes after a pass:", "Once you pass a test the result is permanent — you can't lose the badge by attempting it again.")

h2("Courses and training")
p("Courses are short skill-builders that earn you a badge on your resume when you finish.")
steps([
    "From Profile, tap Training & courses.",
    "Browse the catalogue. Each card shows the duration, language, and what you'll learn.",
    "Tap a course to see the lessons, then tap Enrol.",
    "Work through the lessons at your own pace. Doondo tracks your progress.",
    "Finish all lessons to earn the course badge.",
])

h2("Available now beacon")
p("Workers in the Today mode often want to broadcast \"I'm free right now, hire me.\" The available-now beacon does exactly that.")
steps([
    "On Home, tap I'm available now.",
    "Confirm the trades you're available for and how far you're willing to travel.",
    "Tap Go live. You're now visible to employers nearby for 4 hours.",
])
p("Your beacon expires automatically after 4 hours so you don't end up showing as available a day later. You can extend it from the same toggle, or end it early with End availability.")

h2("Referrals — share and earn")
p("Refer Doondo to friends. When a friend you referred is hired by an employer, Doondo pays you a referral bonus.")
bullets([
    "Earn ₹100 each time a worker you referred is hired.",
    "There's no cap — refer as many friends as you want.",
])
h3("How to refer")
steps([
    "Open any job detail screen and tap the share icon.",
    "Pick how to share — WhatsApp, SMS, or any other app. Doondo attaches your personal referral link automatically.",
    "When your friend taps the link, signs up, and gets hired, the ₹100 is added to your Doondo wallet.",
])
h3("Tracking your referrals")
p("From Profile, tap Referral credit. You'll see:")
bullets([
    "How many people you've referred",
    "How many got hired",
    "Total earned",
])

h2("Group hire — apply with teammates")
p("Some jobs need more than one person at once — building a wall, catering an event, a moving crew. Doondo lets you apply as a team.")
steps([
    "Open a job that says Group hire — looking for N people.",
    "Tap Apply Now.",
    "In the Teammates section, add the name and phone number of each person joining you. Up to 4 teammates.",
    "Submit. The employer sees your whole team as one application.",
])
tip("Your teammates don't need a Doondo account:", "Phone numbers are enough — the employer can call any of them directly.")

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# PART 3 — EMPLOYERS
# ════════════════════════════════════════════════════════════════════════════

h1("Part 3 — For employers")

p("The employer side has four tabs along the bottom: Posts, Applicants, Chat, and Profile. The flow is built for speed — you should be able to post a job in under two minutes and start hearing back the same day.")

h2("Your job posts")
p("The Posts tab lists every job you've put up. Each row shows:")
bullets([
    "Job title and the date you posted it",
    "Number of views",
    "Number of applicants",
    "Status — Open, Paused, or Closed",
])
p("Tap any post to open it. From there you can:")
bullets([
    ("Edit", "Update the title, description, pay, or location."),
    ("Pause", "Stop receiving new applications without deleting the post."),
    ("Close", "Mark the role as filled. Applicants are notified."),
    ("Repost", "Re-open a closed post if you need to hire again."),
    ("Boost", "Push the post higher in seeker feeds for 24 hours (paid)."),
])

h2("Posting a new job")
p("From the Posts tab, tap + New job (the blue floating button).")
steps([
    "Pick the trade — start typing and Doondo suggests matches.",
    "Enter the job title (e.g., \"Cook for small canteen\").",
    "Write a short description. You can also tap the microphone to record an audio description in your own language — seekers love this.",
    "Set the pay. Pick whether it's per day, per shift, per week, per month, or per project. Enter a range (minimum and maximum).",
    "Set the location. Doondo uses your saved business location by default, or you can tap to pick a different point on the map.",
    "Choose job type (full-time, part-time, daily, contract) and work style (onsite, hybrid, remote).",
    "If this is a group hire, set the team size — the number of people you need to hire together.",
    "Mark the post Urgent if you need someone today; it surfaces in the Today tab and gets more views.",
    "Add up to 6 photos of the worksite or the work to be done. Photos triple your response rate.",
    "Tap Publish. The post is live immediately.",
])
tip("Audio descriptions:", "Many blue-collar workers prefer to hear a job in their own language rather than read it. Recording a 20-second audio description in Hindi, Kannada, Tamil, or Telugu can dramatically increase applications.")

h2("Reviewing applicants")
p("The Applicants tab shows every person who's applied across all your posts, grouped by job. Tap a job to see just its applicants.")
p("Each applicant card shows:")
bullets([
    "Photo, name, and verified badge",
    "Rating (if they have one) and number of past hires",
    "Distance from the job location",
    "Skills they've listed",
    "Skill test badges they've passed",
    "When they applied",
])
p("You can sort by:")
bullets([
    ("Best match", "Doondo's recommendation, weighing distance, rating, and skill fit."),
    ("Distance", "Closest first."),
    ("Rating", "Highest-rated first."),
    ("Newest", "Most recent applications first."),
])

h2("Applicant details")
p("Tap any applicant to see the full profile.")
p("You'll see:")
bullets([
    "Their full resume — work history, photos, education",
    "All ratings they've received and the reviews",
    "Skill test results with scores",
    "Their cover letter for this job (if they wrote one)",
    "Teammates (if they applied as a group hire)",
])
p("From the applicant detail screen you can:")
bullets([
    ("Shortlist", "Move them to your shortlist — the seeker is notified."),
    ("Message", "Open a chat. Doondo creates the conversation thread on the seeker side."),
    ("Hire", "Mark them as hired. Doondo notifies the seeker, closes their application, and unlocks photo verification + rating once the work is done."),
    ("Decline", "Politely let them know they didn't get this one. Doondo sends a neutral notification."),
])

h2("Hiring and rating workers")
p("After a hire, two things become possible:")
bullets([
    ("Verify their work photos", "On the applicant detail screen, each of the worker's work photos has a Verify this work overlay. Tap it to confirm they really did this job for you. A green ✓ Verified · N badge appears on the photo — and on the worker's resume forever."),
    ("Leave a rating", "After the work wraps, tap Leave rating on the applicant detail screen. Pick 1–5 stars and optionally add a short comment. Be honest — your rating is shown publicly on the worker's profile."),
])
tip("Why verify photos:", "Workers often add photos of past work to their resume. Your verification — backed by a real hire — turns those photos into trust signals other employers can rely on. It costs you 5 seconds and helps the whole community.")

h2("Chat with workers")
p("The Chat tab works the same way as it does for seekers — text, voice notes, photos, and videos in any conversation. Tap a thread to open it.")
p("Most chats start the moment you shortlist or hire someone. You can also start a chat directly from an applicant's profile.")

h2("Your employer profile")
p("Your profile is what workers see when they look at your jobs. Keep it filled in — verified employers get noticeably more applications.")
bullets([
    ("Business name and logo", "What workers see on your job posts."),
    ("Verified badge", "Earned by completing phone and (for businesses) GST or business registration check."),
    ("About this employer", "A short description of your business, what you hire for, and how you treat workers."),
    ("Rating", "Average stars from workers you've hired. Yes — workers rate you back."),
    ("Location", "Your default worksite. Used to default new job posts."),
    ("Photos", "Up to 6 photos of your workplace. Big trust signal."),
])
p("Tap Edit profile to update any of these. Tap your photo grid to add or remove photos.")

h2("Available workers (live nearby)")
p("Available workers shows you a live list of seekers near you who have turned on the available now beacon — workers who are ready to be hired right this minute.")
steps([
    "From Posts, tap Available now.",
    "You'll see a list of seekers within 10 km, sorted by distance.",
    "Each card shows the trades they're available for, their rating, and how long they've been available.",
    "Tap any worker to open their profile.",
    "From there, tap Hire on the spot or Message to chat first.",
])
tip("Pairs perfectly with Today posts:", "If you posted an urgent job today and want to fill it fast, the Available now list is your best shortcut — these workers are actively waiting for a call.")

h2("Verifying a worker's photos")
p("Covered above under Hiring and rating workers. Quick recap: after a hire, every photo on the worker's profile gets a Verify this work tap target. Tapping it adds your verification to that specific photo and shows on their resume forever as ✓ Verified by employer.")

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# PART 4 — HELP & REFERENCE
# ════════════════════════════════════════════════════════════════════════════

h1("Part 4 — Help and reference")

h2("Staying safe")
bullets([
    "Always meet for the first time in a public place or at the registered worksite.",
    "Never share your bank details, OTP codes, or ID numbers in chat.",
    "Don't pay a fee to apply. Doondo never charges workers to apply for a job. If anyone asks you to pay, report them in the chat by tapping the flag icon.",
    "Set up Safety SOS before your first job (Settings → Safety) so help is one tap away.",
    "If something feels off, leave the worksite. Your safety is more important than the day's pay.",
])

h2("Switching languages")
p("Doondo speaks five languages: English, Hindi, Kannada, Tamil, and Telugu. To switch:")
steps([
    "Open Profile → Settings.",
    "Tap your preferred language under Language.",
    "The whole app re-renders in the new language immediately.",
])
p("Your choice is saved on the device — the next time you open Doondo, it starts in the same language.")

h2("Account help")
h3("Change your phone number")
p("Open Profile → Settings → Account → Email/phone, tap your current phone, enter the new number, and verify with the code Doondo sends.")
h3("Change your password")
p("There's currently no in-app change-password screen — use Forgot password from the sign-in screen to set a new one.")
h3("Delete your account")
p("Settings → Account → Delete account. The in-app flow is being finished — for now, Doondo routes the request to support@doondo.app for manual processing within 48 hours. Deletion erases your applications, chats, and ratings.")

h2("Common questions")
h3("Do I need internet for everything?")
p("Most features need an active connection. Saved jobs (Download Center) and your earnings diary work offline.")
h3("Why is my job not showing for nearby workers?")
p("Three usual causes — make sure (1) your job's location is set correctly, (2) the post isn't Paused, and (3) the pay isn't unrealistically low for your area. Adding photos and an audio description helps too.")
h3("How do I know an employer is real?")
p("Look for the blue verified badge on the employer's profile. Also check their rating and reviews from past workers. If something feels off, walk away.")
h3("How does the ₹100 referral bonus work?")
p("When a friend who signs up through your share link is hired by any employer, Doondo credits ₹100 to your wallet. You can refer as many friends as you want. Bonuses appear in Profile → My Earnings within a day of the hire.")
h3("Why did I get logged out?")
p("Doondo signs you out automatically if your session expires, if you sign in on another device, or if your account is flagged. Sign in again with your password — if you can't, use Forgot password.")

h2("Contacting support")
p("If you're stuck, hit a bug, or need help with an account issue:")
bullets([
    "Email support@doondo.app from the address on your account so the team can identify you.",
    "Include a screenshot if you can — it speeds up the reply.",
    "Most replies come back within 24 hours on business days.",
])

# Footer note
p("")
section_break()
p("")
foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = foot.add_run("Doondo · User Manual · v1.0")
r.italic = True
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

# ─── Save ────────────────────────────────────────────────────────────────────

doc.save(OUT)

# Patch settings.xml — python-docx writes <w:zoom/> without the required
# w:percent attribute, which trips strict schema validators (and Word
# silently ignores it). Re-zip the file with the zoom element fixed.
import zipfile, shutil, tempfile, os as _os
_tmp_dir = tempfile.mkdtemp()
with zipfile.ZipFile(OUT, "r") as zin:
    zin.extractall(_tmp_dir)
_settings_path = _os.path.join(_tmp_dir, "word", "settings.xml")
if _os.path.exists(_settings_path):
    with open(_settings_path, "r", encoding="utf-8") as f:
        _settings = f.read()
    import re as _re
    if "<w:zoom" in _settings and 'w:percent' not in _settings:
        # Add w:percent="100" to whatever <w:zoom .../> form exists.
        _settings = _re.sub(
            r'<w:zoom([^/>]*?)/>',
            r'<w:zoom\1 w:percent="100"/>',
            _settings,
        )
        with open(_settings_path, "w", encoding="utf-8") as f:
            f.write(_settings)
# Repack
_tmp_out = OUT + ".tmp"
with zipfile.ZipFile(_tmp_out, "w", zipfile.ZIP_DEFLATED) as zout:
    for root, _, files in _os.walk(_tmp_dir):
        for fn in files:
            full = _os.path.join(root, fn)
            arc = _os.path.relpath(full, _tmp_dir)
            zout.write(full, arc)
shutil.move(_tmp_out, OUT)
shutil.rmtree(_tmp_dir)

print(f"Wrote {OUT}")
